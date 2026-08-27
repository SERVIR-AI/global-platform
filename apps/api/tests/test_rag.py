"""Workstream E slices E1–E3: docloader, corpus store + retrieval, citation seam.

No network — embeddings come from a deterministic hash-based stub (identical
text -> identical vector, different text -> ~orthogonal), so similarity scores
are controllable without an API key. Chunking, persistence, filtering, and the
routes stay real. The real-embeddings round trip lives in test_rag_live.py
(opt-in, RAG_LIVE=1).
"""

import hashlib
import json
from io import BytesIO

import numpy as np
import pytest
from fastapi.testclient import TestClient

from app.config import get_settings
from app.main import app
from app.rag import docloader
from app.rag.embed import ProviderEmbedder
from app.rag.store import Corpus, CorpusError, _chunk, source_block


class HashEmbedder:
    """Deterministic unit vectors seeded by a text hash — no network."""

    calls = 0

    def embed(self, texts):
        HashEmbedder.calls += 1
        out = []
        for t in texts:
            seed = int.from_bytes(hashlib.sha1(t.encode()).digest()[:8], "big")
            v = np.random.default_rng(seed).standard_normal(64).astype(np.float32)
            out.append(v / np.linalg.norm(v))
        return np.stack(out)


@pytest.fixture
def rag_env(monkeypatch, tmp_path):
    """Isolated cache dir + stubbed provider embeddings (for route-built Corpora).
    Pins the relevance floor too, so an ambient RAG_MIN_RELEVANCE env/.env value
    can't change what these hermetic tests assert."""
    monkeypatch.setattr(get_settings(), "cache_dir", tmp_path)
    monkeypatch.setattr(get_settings(), "rag_min_relevance", 0.5)
    monkeypatch.setattr(ProviderEmbedder, "embed",
                        lambda self, texts: HashEmbedder().embed(texts))
    return tmp_path


# --- E1: docloader -----------------------------------------------------------

def test_docloader_text_formats():
    assert "El Niño" in docloader.extract_text("El Niño rainfall outlook".encode(), "a.txt")
    assert "maize" in docloader.extract_text(b"# Outlook\n\nmaize conditions", "a.md")


def test_docloader_docx_keeps_tables():
    """Bulletins put the numbers in tables — extraction must keep them."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Seasonal outlook for the Rift Valley.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "maize production"
    table.rows[0].cells[1].text = "2.1m tonnes"
    buf = BytesIO()
    doc.save(buf)
    text = docloader.extract_text(buf.getvalue(), "bulletin.docx")
    assert "Seasonal outlook" in text
    assert "maize production | 2.1m tonnes" in text


def test_docloader_image_only_pdf_fails_loudly():
    """A PDF with no extractable text (scanned/needs OCR) must be a visible gap,
    never a silently-empty phantom document."""
    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = BytesIO()
    writer.write(buf)
    with pytest.raises(docloader.EmptyDocument, match="OCR"):
        docloader.extract_text(buf.getvalue(), "scan.pdf")


def test_docloader_rejects_garbage_and_unknown_formats():
    with pytest.raises(docloader.DocLoadError, match="unreadable PDF"):
        docloader.extract_text(b"not a pdf at all", "broken.pdf")
    with pytest.raises(docloader.DocLoadError, match=".txt, .md, .pdf, .docx"):
        docloader.extract_text(b"data", "image.png")
    with pytest.raises(docloader.EmptyDocument):
        docloader.extract_text(b"   ", "empty.txt")


def _text_pdf(text: str) -> bytes:
    """A minimal but structurally valid one-page PDF with a real text object and
    a correct xref table — the hermetic happy path for the primary bulletin format."""
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = b"%PDF-1.4\n"
    offsets = []
    for i, obj in enumerate(objs, 1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    out += b"".join(f"{off:010d} 00000 n \n".encode() for off in offsets)
    return out + (b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n"
                  + str(xref).encode() + b"\n%%EOF\n")


def test_docloader_pdf_happy_path_is_hermetic():
    text = docloader.extract_text(
        _text_pdf("El Nino rainfall outlook for Kenya maize."), "bulletin.pdf")
    assert "El Nino rainfall outlook for Kenya maize." in text


# --- E2: chunking ------------------------------------------------------------

def test_chunker_packs_paragraphs_and_splits_oversized_ones():
    paras = [f"Paragraph {i} about crop conditions." for i in range(30)]
    text = "\n\n".join(paras) + "\n\n" + ("A very long sentence. " * 200)
    chunks = _chunk(text, target=500)
    assert all(len(c) <= 510 for c in chunks)          # target honored (small slack)
    joined = " ".join(chunks)
    assert all(p in joined for p in paras)             # nothing dropped
    assert len(chunks) > 5                             # actually split, not one blob


# --- E2: corpus ingest + retrieval --------------------------------------------

KENYA_DOC = ("El Niño seasons historically bring above-normal rainfall to the Kenyan "
             "highlands, favouring maize yields in the long rains.\n\n"
             "In 2015-16, maize production in the Rift Valley rose despite localized "
             "flooding in low-lying areas.")
ZAMBIA_DOC = ("El Niño years are associated with rainfall deficits across Zambia, "
              "with maize the most affected staple crop.")

KENYA_META = {"source": "FEWS NET", "title": "Kenya analog-year review", "pub_date": "2024-03",
              "countries": ["Kenya"], "crops": ["maize"], "temporal": "retrospective",
              "validation": "single-agency"}
ZAMBIA_META = {"source": "FAO GIEWS", "title": "Zambia outlook", "pub_date": "2026-06",
               "countries": ["Zambia"], "crops": ["maize"], "temporal": "forecast",
               "validation": "single-agency"}


@pytest.fixture
def corpus(rag_env):
    c = Corpus("test", embedder=HashEmbedder())
    c.ingest(KENYA_DOC, KENYA_META)
    c.ingest(ZAMBIA_DOC, ZAMBIA_META)
    return c


def test_search_returns_scored_hits_with_provenance(corpus):
    """An exact-text query scores 1.0 against its own chunk and carries the
    metadata the citation layer needs."""
    hits = corpus.search(ZAMBIA_DOC, k=3)
    assert hits and hits[0]["score"] == pytest.approx(1.0)
    assert hits[0]["metadata"]["source"] == "FAO GIEWS"
    assert hits[0]["doc_id"] and hits[0]["id"].startswith(hits[0]["doc_id"])


def test_below_threshold_returns_nothing(corpus):
    """The decline contract: irrelevant queries return [], never the least-bad match."""
    assert corpus.search("submarine navigation protocols") == []


def test_filters_narrow_before_similarity(corpus):
    """A perfect-similarity chunk from the wrong country must not survive the filter."""
    assert corpus.search(ZAMBIA_DOC, countries="Kenya") == []
    hits = corpus.search(ZAMBIA_DOC, countries="Zambia", temporal="forecast")
    assert hits and hits[0]["metadata"]["countries"] == ["Zambia"]


def test_filters_run_before_topk_not_after(rag_env):
    """Falsifies the ordering: three higher-scoring out-of-filter chunks would fill
    top-k first if filtering ran after ranking — the filtered-in hit must survive."""
    vectors = {"query": [1.0, 0.0], "kenya a": [0.99, 0.141], "kenya b": [0.98, 0.199],
               "kenya c": [0.97, 0.243], "zambia doc": [0.9, 0.436]}

    class DictEmbedder:
        def embed(self, texts):
            arr = np.array([vectors[t] for t in texts], dtype=np.float32)
            return arr / np.linalg.norm(arr, axis=1, keepdims=True)

    c = Corpus("filter-order", embedder=DictEmbedder())
    for t in ("kenya a", "kenya b", "kenya c"):
        c.ingest(t, {"source": "s", "countries": ["Kenya"]})
    c.ingest("zambia doc", {"source": "s", "countries": ["Zambia"]})
    hits = c.search("query", k=2, countries="Zambia")
    assert [h["text"] for h in hits] == ["zambia doc"]


def test_list_valued_filters_match_on_intersection(corpus):
    """countries=["Zambia", "Malawi"] must match a chunk tagged countries=["Zambia"]."""
    hits = corpus.search(ZAMBIA_DOC, countries=["Zambia", "Malawi"])
    assert hits and hits[0]["metadata"]["countries"] == ["Zambia"]
    assert corpus.search(ZAMBIA_DOC, countries=["Malawi", "Uganda"]) == []


def test_reingest_is_idempotent(corpus):
    """Same text twice -> no new chunks, no new embedding calls."""
    before_calls, before_chunks = HashEmbedder.calls, len(corpus._chunks)
    res = corpus.ingest(KENYA_DOC, KENYA_META)
    assert res["already_ingested"] is True
    assert "metadata_updated" not in res
    assert HashEmbedder.calls == before_calls
    assert len(corpus._chunks) == before_chunks


def test_reingest_with_new_metadata_updates_provenance(corpus):
    """Corrected metadata on the same text must take effect and say so — silently
    keeping stale provenance would be an honesty bug (no re-embedding needed)."""
    before_calls = HashEmbedder.calls
    fixed = {**KENYA_META, "validation": "multi-agency-consensus"}
    res = corpus.ingest(KENYA_DOC, fixed)
    assert res["already_ingested"] is True and res["metadata_updated"] is True
    assert HashEmbedder.calls == before_calls          # nothing re-embedded
    assert corpus.search(KENYA_DOC)[0]["metadata"]["validation"] == "multi-agency-consensus"


def test_empty_text_is_a_typed_error_not_a_crash(corpus):
    """Whitespace-only text must raise the same typed error the docloader uses,
    never reach the embedder (or a numpy stack trace)."""
    with pytest.raises(docloader.EmptyDocument):
        corpus.ingest("   \n\n  ", {"source": "s"})


def test_embedder_model_mismatch_fails_loudly(corpus, rag_env):
    """A corpus embedded with one model must refuse to serve another — mixed
    vector spaces produce confidently-wrong scores, the worst failure mode."""
    class OtherEmbedder(HashEmbedder):
        provider, model = "openai", "text-embedding-3-small"
    with pytest.raises(CorpusError, match="embedded with"):
        Corpus("test", embedder=OtherEmbedder())


def test_ingest_archives_the_original_and_backfills_on_reingest(rag_env):
    """The exact ingested bytes stay auditable: archived at ingest, re-archived on
    a duplicate ingest if the archive went missing."""
    c = Corpus("audit", embedder=HashEmbedder())
    res = c.ingest("El Nino outlook text for the audit trail.", {"source": "s"},
                   raw=b"%PDF-original-bytes", filename="brief.pdf")
    path = c.raw_path(res["doc_id"])
    assert path and path.read_bytes() == b"%PDF-original-bytes" and path.suffix == ".pdf"
    path.unlink()
    c.ingest("El Nino outlook text for the audit trail.", {"source": "s"},
             raw=b"%PDF-original-bytes", filename="brief.pdf")
    assert c.raw_path(res["doc_id"]).read_bytes() == b"%PDF-original-bytes"


def test_corpus_persists_and_reloads(corpus, rag_env):
    """A fresh Corpus instance answers from disk — the library outlives the process."""
    reloaded = Corpus("test", embedder=HashEmbedder())
    assert len(reloaded.documents()) == 2
    assert reloaded.search(KENYA_DOC)[0]["metadata"]["title"] == "Kenya analog-year review"


def test_torn_corpus_fails_loudly(corpus, rag_env):
    """chunks/vectors disagreeing is corruption — a clear error naming the fix,
    never silently answering from misaligned rows."""
    jl = rag_env / "rag" / "test" / "chunks.jsonl"
    lines = jl.read_text().splitlines()
    jl.write_text("\n".join(lines[:-1]) + "\n")        # drop one chunk row
    with pytest.raises(CorpusError, match="delete .* and re-ingest"):
        Corpus("test", embedder=HashEmbedder())


def test_half_present_corpus_is_corruption_not_emptiness(corpus, rag_env):
    """One of the corpus files missing must be a loud error: loading it as an
    empty library would mask corruption as absence — and the next ingest would
    then overwrite the surviving file, destroying the library."""
    (rag_env / "rag" / "test" / "embeddings.npy").unlink()
    with pytest.raises(CorpusError, match="torn .*embeddings.npy missing"):
        Corpus("test", embedder=HashEmbedder())


def test_parallel_ingests_all_survive(rag_env):
    """Concurrent ingests serialize on the corpus lock: every document survives,
    nothing lost to a last-writer-wins race, and the corpus stays consistent."""
    from concurrent.futures import ThreadPoolExecutor
    docs = [f"Document number {i} about crop conditions in region {i}." for i in range(8)]

    def ingest(text):
        return Corpus("race", embedder=HashEmbedder()).ingest(text, {"source": "s"})
    with ThreadPoolExecutor(max_workers=8) as pool:
        results = list(pool.map(ingest, docs))
    assert all(r["already_ingested"] is False for r in results)
    final = Corpus("race", embedder=HashEmbedder())
    assert len(final.documents()) == 8                 # no lost updates


# --- E3: the citation seam -----------------------------------------------------

def test_source_block_formats_citations(corpus):
    hits = corpus.search(ZAMBIA_DOC, k=1)
    block = source_block(hits)
    assert block.startswith("[1] FAO GIEWS — Zambia outlook — 2026-06 — validation: single-agency")
    assert "rainfall deficits across Zambia" in block


def test_source_block_includes_the_source_url(rag_env):
    c = Corpus("cited", embedder=HashEmbedder())
    c.ingest("Some bulletin passage.", {"source": "FAO", "url": "https://fao.org/x.pdf"})
    block = source_block(c.search("Some bulletin passage.", k=1))
    assert "https://fao.org/x.pdf" in block


def test_source_block_survives_missing_metadata():
    """Sparse metadata renders honestly ('unknown source'), never raises."""
    block = source_block([{"metadata": {}, "text": "some passage"}])
    assert block == "[1] unknown source\nsome passage"


# --- the endpoints -------------------------------------------------------------

class FakeResp:
    def __init__(self, status=200, content=b"", headers=None):
        self.status_code = status
        self._content = content
        self.headers = headers or {}

    def raise_for_status(self):
        return None

    def iter_content(self, size):
        for i in range(0, len(self._content), size):
            yield self._content[i:i + size]


@pytest.fixture
def fetch_env(rag_env, monkeypatch):
    """Bypass the SSRF host check (tests must not need DNS) and script the fetch."""
    from app.food_security import routes as fs_routes
    monkeypatch.setattr(fs_routes, "_assert_public_http", lambda url: None)

    def scripted(responses):
        it = iter(responses)

        def fake_get(url, **kw):
            nxt = next(it)
            if isinstance(nxt, Exception):
                raise nxt
            return nxt
        monkeypatch.setattr(fs_routes.requests, "get", fake_get)
    return scripted


META = {"source": "FEWS NET", "title": "t", "pub_date": "2026-06"}


def test_ingest_and_search_round_trip_over_http(rag_env):
    client = TestClient(app)
    r = client.post("/api/food-security/rag/ingest",
                    json={"text": KENYA_DOC, "metadata": KENYA_META})
    assert r.status_code == 200
    assert r.json()["chunks"] >= 1 and r.json()["already_ingested"] is False

    r = client.get("/api/food-security/rag/search", params={"q": KENYA_DOC, "k": 2})
    body = r.json()
    assert r.status_code == 200
    assert body["hits"][0]["metadata"]["source"] == "FEWS NET"
    assert body["min_relevance"] == get_settings().rag_min_relevance
    assert "note" not in body


def test_every_insight_traces_back_to_its_source(rag_env):
    """The trust contract end-to-end: a hit carries the live source URL AND our
    archived copy of the exact bytes we read; the catalog lists the whole boundary."""
    client = TestClient(app)
    client.post("/api/food-security/rag/ingest",
                json={"text": KENYA_DOC,
                      "metadata": {**KENYA_META, "url": "https://fews.net/ken.pdf"}})
    hit = client.get("/api/food-security/rag/search",
                     params={"q": KENYA_DOC, "k": 1}).json()["hits"][0]
    assert hit["trace"]["source_url"] == "https://fews.net/ken.pdf"
    archived = hit["trace"]["archived_copy"]
    assert archived and hit["doc_id"] in archived
    r = client.get(archived)
    assert r.status_code == 200 and KENYA_DOC.encode() in r.content

    catalog = client.get("/api/food-security/rag/documents").json()
    assert catalog["documents"] == 1
    assert catalog["items"][0]["trace"]["archived_copy"] == archived

    assert client.get("/api/food-security/rag/document/0000000000000000").status_code == 404
    assert client.get("/api/food-security/rag/document/not-a-doc-id").status_code == 404


def test_search_decline_notes_name_the_actual_cause(rag_env):
    """Three different empty results, three different truths — never one shrug."""
    client = TestClient(app)
    r = client.get("/api/food-security/rag/search", params={"q": "anything"})
    assert "library is empty" in r.json()["note"]

    client.post("/api/food-security/rag/ingest",
                json={"text": KENYA_DOC, "metadata": KENYA_META})
    r = client.get("/api/food-security/rag/search",
                   params={"q": KENYA_DOC, "country": "Zambia"})
    assert "none match the filters" in r.json()["note"]
    assert "country='Zambia'" in r.json()["note"]

    r = client.get("/api/food-security/rag/search",
                   params={"q": "submarine navigation protocols"})
    assert "below the relevance floor" in r.json()["note"]

    assert client.get("/api/food-security/rag/search",
                      params={"q": "x", "k": 0}).status_code == 422  # k >= 1 enforced


def test_ingest_enforces_provenance(rag_env):
    """No source, no entry: a chunk without provenance can never be cited honestly."""
    client = TestClient(app)
    r = client.post("/api/food-security/rag/ingest", json={"text": KENYA_DOC})
    assert r.status_code == 422 and "provenance" in r.json()["detail"]

    r = client.post("/api/food-security/rag/ingest",
                    json={"text": KENYA_DOC,
                          "metadata": {**META, "temporal": "sometime"}})
    assert r.status_code == 422 and "temporal" in r.json()["detail"]

    r = client.post("/api/food-security/rag/ingest",
                    json={"text": "   \n  ", "metadata": META})
    assert r.status_code == 422 and "no usable text" in r.json()["detail"]

    r = client.post("/api/food-security/rag/ingest", json={"metadata": META})
    assert r.status_code == 400 and "url or text" in r.json()["detail"]


def test_ingest_blocks_non_public_urls_without_network():
    """SSRF guard: scheme allowlist + private/loopback rejection (hosts resolve
    locally, so this needs no DNS)."""
    client = TestClient(app)
    for url in ("ftp://internal/x.pdf", "http://127.0.0.1/x.pdf", "http://localhost/x.pdf"):
        r = client.post("/api/food-security/rag/ingest",
                        json={"url": url, "metadata": META})
        assert r.status_code == 400, url


def test_ingest_caps_document_size(fetch_env, monkeypatch):
    from app.food_security import routes as fs_routes
    client = TestClient(app)
    fetch_env([FakeResp(headers={"Content-Length": str(10 ** 12)})])
    r = client.post("/api/food-security/rag/ingest",
                    json={"url": "https://big.example/x.pdf", "metadata": META})
    assert r.status_code == 413                        # declared size rejected early

    monkeypatch.setattr(fs_routes, "_MAX_DOC_BYTES", 100)
    fetch_env([FakeResp(content=b"x" * 500)])          # header lies; stream capped anyway
    r = client.post("/api/food-security/rag/ingest",
                    json={"url": "https://liar.example/x.pdf", "metadata": META})
    assert r.status_code == 413


def test_ingest_follows_validated_redirects_and_handles_failures(fetch_env):
    """cropmonitor.org serves bulletins via a 302 — followed, with every hop
    re-validated; a dead host is a clean 502, not a stack trace."""
    client = TestClient(app)
    fetch_env([FakeResp(status=302, headers={"Location": "https://cdn.example/real.md"}),
               FakeResp(content=b"# Bulletin\n\nEl Nino rainfall outlook for Kenya.")])
    r = client.post("/api/food-security/rag/ingest",
                    json={"url": "https://portal.example/doc", "filename": "real.md",
                          "metadata": META})
    assert r.status_code == 200 and r.json()["chunks"] == 1

    import requests as requests_lib
    fetch_env([requests_lib.ConnectionError("boom")])
    r = client.post("/api/food-security/rag/ingest",
                    json={"url": "https://down.example/x.pdf", "metadata": META})
    assert r.status_code == 502 and "could not fetch" in r.json()["detail"]


def test_ingest_rejects_unextractable_documents(fetch_env):
    """A fetched document the loader can't turn into text is a typed 422, not a
    phantom ingest."""
    fetch_env([FakeResp(content=b"\x89PNG...")])
    r = TestClient(app).post("/api/food-security/rag/ingest",
                             json={"url": "https://example.org/chart.png",
                                   "metadata": META})
    assert r.status_code == 422 and "supported" in r.json()["detail"]
