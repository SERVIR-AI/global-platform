"""E1 — document bytes -> clean text, or a typed DocLoadError (never a silent
empty ingest: an image-only PDF needing OCR must be a visible gap)."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class DocLoadError(Exception):
    """The document could not be turned into text."""


class EmptyDocument(DocLoadError):
    """No usable text (image-only PDF, empty file). Also raised by store.ingest."""


_MIN_CHARS = 20  # under this, extraction found no real text


def extract_text(data: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in (".txt", ".md"):
        text = data.decode("utf-8", errors="replace")
    elif suffix == ".pdf":
        text = _pdf_text(data, filename)
    elif suffix == ".docx":
        text = _docx_text(data, filename)
    else:
        raise DocLoadError(
            f"Cannot extract text from {filename!r} — supported: .txt, .md, .pdf, .docx")
    if len(text.strip()) < _MIN_CHARS:
        raise EmptyDocument(
            f"{filename!r} yielded no usable text — likely an image-only PDF "
            "(needs OCR) or an empty file")
    return text


def _pdf_text(data: bytes, filename: str) -> str:
    try:
        reader = PdfReader(BytesIO(data))
        return "\n\n".join(filter(None, (page.extract_text() for page in reader.pages)))
    except Exception as exc:  # pypdf raises a zoo of types on malformed files
        raise DocLoadError(f"unreadable PDF {filename!r}: {exc}") from exc


def _docx_text(data: bytes, filename: str) -> str:
    try:
        doc = Document(BytesIO(data))
    except Exception as exc:
        raise DocLoadError(f"unreadable DOCX {filename!r}: {exc}") from exc
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:  # bulletins put the numbers in tables
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)
